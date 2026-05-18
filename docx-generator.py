"""
Multiversum DOCX Generator
==========================
Generates CI-compliant Word documents (.docx) with:
- Multiversum Logo top-right in header (embedded, no network required)
- Dokumentensteckbrief (first-page metadata table with Word Content Controls)
- Änderungshistorie table
- Native Word TOC (auto-generated, F9 to update in Word)
- All heading styles pre-defined (H1–H4)
- Footer with company info + page numbers
- Full Multiversum CI (Arial, #333333, #006DB0, #3C6E89, #F2FF62)

Usage:
    python3 docx-generator.py --out "Mein Dokument.docx" --title "ISMS Richtlinie"
    python3 docx-generator.py --out output.docx --title "Titel" --subtitle "Untertitel" \
        --author "Max Mustermann" --classification "intern"

Or import and use the factory:
    from docx_generator import create_mv_document
    doc = create_mv_document(title="My Doc")
    doc.save("output.docx")
"""

import argparse
import copy
import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

# ── Logo path (bundled with skill, always available) ──────────────────────────
SKILL_DIR = Path(__file__).parent
LOGO_PATH = SKILL_DIR / "assets" / "Logo_MV_MVW.png"

# ── CI Color palette ──────────────────────────────────────────────────────────
C_TEXT    = RGBColor(0x33, 0x33, 0x33)   # #333333 — primary text
C_BLUE    = RGBColor(0x00, 0x6D, 0xB0)   # #006DB0 — Multiversum blue (titles)
C_TEAL    = RGBColor(0x3C, 0x6E, 0x89)   # #3C6E89 — accent / CTAs
C_STEEL   = RGBColor(0x5D, 0x62, 0x69)   # #5D6269 — secondary text
C_SILVER  = RGBColor(0xA4, 0xA7, 0xAB)   # #A4A7AB — muted / footer
C_LGRAY   = RGBColor(0xF5, 0xF5, 0xF3)   # #F5F5F3 — light backgrounds
C_DGRAY   = RGBColor(0xE1, 0xE2, 0xE3)   # #E1E2E3 — dividers
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)   # #FFFFFF
C_YELLOW  = RGBColor(0xF2, 0xFF, 0x62)   # #F2FF62 — Neon Gelb (accent only)

FONT_NAME = "Arial"


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(rgb))
    tcPr.append(shd)


def _set_cell_border(cell, **borders):
    """borders: top/bottom/left/right = (size_pt, color_hex, style)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, (sz, color, val) in borders.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(int(float(sz) * 8)))  # 1/8 pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _apply_font(run, size_pt, bold=False, color: RGBColor = None, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _para_spacing(para, before_pt=0, after_pt=6, line_rule=WD_LINE_SPACING.MULTIPLE, lines=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing_rule = line_rule
    pf.line_spacing = lines


def _add_toc(doc: Document):
    """Insert a native Word TOC field (updates on open or F9)."""
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    run = para.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    # Force-update hint so Word refreshes on first open
    para._p.get_or_add_pPr()
    return para


def _add_content_control(cell, tag: str, placeholder: str, value: str = ""):
    """Add a structured document tag (SDT) content control to a table cell."""
    para = cell.paragraphs[0]
    para.clear()

    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), placeholder)
    sdtPr.append(alias)

    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    sdtPr.append(tag_el)

    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    inner_para = OxmlElement("w:p")
    inner_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")  # 10pt
    rPr.append(sz)
    inner_run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = value if value else placeholder
    inner_run.append(t)
    inner_para.append(inner_run)
    sdtContent.append(inner_para)
    sdt.append(sdtContent)
    para._p.append(sdt)


# ─────────────────────────────────────────────────────────────────────────────
# Style definitions
# ─────────────────────────────────────────────────────────────────────────────

def _define_styles(doc: Document):
    """Configure all named paragraph styles to match Multiversum CI."""
    styles = doc.styles

    def _get_or_create(name, base="Normal"):
        try:
            return styles[name]
        except KeyError:
            return styles.add_style(name, 1)  # WD_STYLE_TYPE.PARAGRAPH = 1

    # ── Normal ──
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal.font.color.rgb = C_TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # ── Heading 1 ──
    h1 = styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = C_TEXT
    h1.font.italic = False
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    # Bottom border line in teal
    pPr = h1._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), _hex(C_TEAL))
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Heading 2 ──
    h2 = styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = C_TEXT
    h2.font.italic = False
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)

    # ── Heading 3 ──
    h3 = styles["Heading 3"]
    h3.font.name = FONT_NAME
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = C_STEEL
    h3.font.italic = False
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)

    # ── Heading 4 ──
    h4 = styles["Heading 4"]
    h4.font.name = FONT_NAME
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = C_STEEL
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(2)

    # ── Body Text ──
    body = _get_or_create("MV Body")
    body.base_style = styles["Normal"]
    body.font.name = FONT_NAME
    body.font.size = Pt(11)
    body.font.color.rgb = C_TEXT
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.5

    # ── Caption ──
    try:
        cap = styles["Caption"]
    except KeyError:
        cap = styles.add_style("Caption", 1)
    cap.font.name = FONT_NAME
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = C_STEEL

    # ── MV Label (for Steckbrief table labels) ──
    lbl = _get_or_create("MV Label")
    lbl.font.name = FONT_NAME
    lbl.font.size = Pt(10)
    lbl.font.bold = True
    lbl.font.color.rgb = C_BLUE

    # ── MV Value (for Steckbrief table values) ──
    val = _get_or_create("MV Value")
    val.font.name = FONT_NAME
    val.font.size = Pt(10)
    val.font.color.rgb = C_TEXT

    # ── TOC styles ──
    for level in range(1, 4):
        name = f"TOC {level}"
        try:
            toc_s = styles[name]
        except KeyError:
            toc_s = styles.add_style(name, 1)
        toc_s.font.name = FONT_NAME
        toc_s.font.size = Pt(11 - level)
        toc_s.font.bold = (level == 1)
        toc_s.font.color.rgb = C_TEXT
        toc_s.paragraph_format.space_before = Pt(2 if level > 1 else 4)
        toc_s.paragraph_format.space_after = Pt(2)
        indent = Cm(0.5 * (level - 1))
        toc_s.paragraph_format.left_indent = indent


# ─────────────────────────────────────────────────────────────────────────────
# Header: Logo top-right
# ─────────────────────────────────────────────────────────────────────────────

def _build_header(section, doc_title: str = ""):
    header = section.header
    # Clear default paragraph
    for para in header.paragraphs:
        para.clear()

    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Optional left-side document title in light gray
    if doc_title:
        left_run = hp.add_run(doc_title + "\t")
        _apply_font(left_run, 9, color=C_SILVER)
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Use tab stop to push logo right
        pPr = hp._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9026")  # ~16cm = right margin
        tabs.append(tab)
        pPr.append(tabs)

    logo_run = hp.add_run()
    if LOGO_PATH.exists():
        logo_run.add_picture(str(LOGO_PATH), height=Cm(0.9))
    else:
        # Fallback text logo
        logo_run.text = "MULTIVERSUM"
        _apply_font(logo_run, 11, bold=True, color=C_TEXT)

    # Bottom border on header
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), _hex(C_DGRAY))
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# Footer: company + page number
# ─────────────────────────────────────────────────────────────────────────────

def _build_footer(section, classification: str = "intern"):
    footer = section.footer
    for para in footer.paragraphs:
        para.clear()

    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)

    # Top border
    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), _hex(C_DGRAY))
    pBdr.append(top)
    pPr.append(pBdr)

    # Tab stops: center + right
    tabs = OxmlElement("w:tabs")
    for pos, val in [("4680", "center"), ("9360", "right")]:
        t = OxmlElement("w:tab")
        t.set(qn("w:val"), val)
        t.set(qn("w:pos"), pos)
        tabs.append(t)
    pPr.append(tabs)

    # Left: company + classification
    left = fp.add_run(f"Multiversum GmbH · Hamburg")
    _apply_font(left, 9, color=C_SILVER)

    # Center: classification marker
    cls_run = fp.add_run(f"\t[{classification.upper()}]")
    _apply_font(cls_run, 9, color=C_SILVER)

    # Right: page / of pages
    fp.add_run("\t")
    page_run = fp.add_run()
    _apply_font(page_run, 9, color=C_SILVER)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    page_run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    page_run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld2)

    sep = fp.add_run(" / ")
    _apply_font(sep, 9, color=C_SILVER)

    total_run = fp.add_run()
    _apply_font(total_run, 9, color=C_SILVER)
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "begin")
    total_run._r.append(fld3)
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    total_run._r.append(instr2)
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    total_run._r.append(fld4)


# ─────────────────────────────────────────────────────────────────────────────
# Dokumentensteckbrief (first-page metadata block)
# ─────────────────────────────────────────────────────────────────────────────

def _build_steckbrief(doc: Document, meta: dict):
    """
    Inserts a Dokumentensteckbrief table with Word Content Controls.
    meta keys: title, subtitle, classification, valid_from, responsible,
               approved_by, approved_on, storage
    """
    # Section title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Dokumentensteckbrief")
    _apply_font(title_run, 11, bold=True, color=C_BLUE)
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(4)

    fields = [
        ("Dateiname",                "mv_filename",       "steckbrief_filename",    meta.get("title", "")),
        ("Klassifizierung",          "mv_classification", "steckbrief_class",
         meta.get("classification", "öffentlich / intern / vertraulich / streng vertraulich")),
        ("Gültig ab",                "mv_valid_from",     "steckbrief_valid",       meta.get("valid_from", "")),
        ("Dokumentenverantwortlich", "mv_responsible",    "steckbrief_responsible", meta.get("responsible", "")),
        ("Freigegeben von",          "mv_approved_by",    "steckbrief_approved_by", meta.get("approved_by", "")),
        ("Freigegeben am",           "mv_approved_on",    "steckbrief_approved_on", meta.get("approved_on", "")),
        ("Ablageort",                "mv_storage",        "steckbrief_storage",     meta.get("storage", "")),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths: label 4.3cm, value 12.4cm
    for row_idx, (label, cc_alias, cc_tag, value) in enumerate(fields):
        row = table.rows[row_idx]
        row.cells[0].width = Cm(4.3)
        row.cells[1].width = Cm(12.4)

        # Label cell
        lc = row.cells[0]
        _set_cell_bg(lc, C_WHITE)
        _set_cell_border(lc,
            top=("0.5", _hex(C_DGRAY), "single"),
            bottom=("0.5", _hex(C_DGRAY), "single"),
            left=("0.5", _hex(C_DGRAY), "single"),
            right=("0.5", _hex(C_DGRAY), "single"),
        )
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        _apply_font(lr, 10, bold=True, color=C_BLUE)
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(2)

        # Value cell with content control
        vc = row.cells[1]
        _set_cell_bg(vc, C_WHITE)
        _set_cell_border(vc,
            top=("0.5", _hex(C_DGRAY), "single"),
            bottom=("0.5", _hex(C_DGRAY), "single"),
            left=("0.5", _hex(C_DGRAY), "single"),
            right=("0.5", _hex(C_DGRAY), "single"),
        )
        _add_content_control(vc, cc_tag, cc_alias, value)

    doc.add_paragraph()  # spacing after table


# ─────────────────────────────────────────────────────────────────────────────
# Änderungshistorie
# ─────────────────────────────────────────────────────────────────────────────

def _build_aenderungshistorie(doc: Document, entries: list = None):
    """
    entries: list of dicts with keys: version, date, chapter, description, reviewed_by
    Defaults to one example entry if empty.
    """
    if not entries:
        entries = [
            {"version": "0.1", "date": date.today().strftime("%d.%m.%Y"),
             "chapter": "alle", "description": "Erstellung", "reviewed_by": ""},
        ]

    h = doc.add_paragraph("Änderungshistorie")
    h.style = doc.styles["Heading 2"]

    headers = ["Version", "Datum", "Kapitel", "Änderungsbeschreibung", "Überprüft von"]
    col_widths = [Cm(1.8), Cm(3.2), Cm(2.0), Cm(7.0), Cm(3.0)]

    table = doc.add_table(rows=1 + len(entries), cols=5)
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hrow.cells[i]
        cell.width = w
        _set_cell_bg(cell, C_TEXT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        _apply_font(r, 9, bold=True, color=C_WHITE)

    # Data rows
    for r_idx, entry in enumerate(entries):
        row = table.rows[r_idx + 1]
        values = [
            entry.get("version", ""),
            entry.get("date", ""),
            entry.get("chapter", ""),
            entry.get("description", ""),
            entry.get("reviewed_by", ""),
        ]
        bg = C_WHITE if r_idx % 2 == 0 else C_LGRAY
        for c_idx, val in enumerate(values):
            cell = row.cells[c_idx]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx < 3 else WD_ALIGN_PARAGRAPH.LEFT
            r_run = p.add_run(val)
            _apply_font(r_run, 9, color=C_TEXT)

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Page setup
# ─────────────────────────────────────────────────────────────────────────────

def _configure_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = False


# ─────────────────────────────────────────────────────────────────────────────
# Main factory
# ─────────────────────────────────────────────────────────────────────────────

def create_mv_document(
    title: str = "Dokumententitel",
    subtitle: str = "",
    author: str = "",
    classification: str = "intern",
    valid_from: str = "",
    responsible: str = "",
    approved_by: str = "",
    approved_on: str = "",
    storage: str = "",
    changelog_entries: list = None,
    include_toc: bool = True,
    include_sample_content: bool = False,
) -> Document:
    """
    Creates and returns a Multiversum CI-compliant Document object.
    Call doc.save("filename.docx") to write to disk.
    """
    doc = Document()
    _define_styles(doc)

    section = doc.sections[0]
    _configure_page(section)
    _build_header(section, title)
    _build_footer(section, classification)

    # ── Cover / Title ──────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    _apply_font(title_run, 24, bold=True, color=C_TEXT)
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(6)

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run(subtitle)
        _apply_font(sub_run, 14, color=C_STEEL)
        sub_para.paragraph_format.space_after = Pt(12)

    # Divider
    div = doc.add_paragraph()
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), _hex(C_TEAL))
    pBdr.append(bot)
    pPr.append(pBdr)
    div.paragraph_format.space_after = Pt(16)

    # ── Dokumentensteckbrief ───────────────────────────────────────────────
    meta = dict(
        title=title,
        classification=classification,
        valid_from=valid_from,
        responsible=responsible,
        approved_by=approved_by,
        approved_on=approved_on,
        storage=storage,
    )
    _build_steckbrief(doc, meta)

    # ── Änderungshistorie ──────────────────────────────────────────────────
    _build_aenderungshistorie(doc, changelog_entries)

    # ── Page break + TOC ──────────────────────────────────────────────────
    doc.add_page_break()

    if include_toc:
        toc_title = doc.add_paragraph("Inhaltsverzeichnis")
        toc_title.style = doc.styles["Heading 1"]

        _add_toc(doc)
        doc.add_paragraph()
        doc.add_page_break()

    # ── Optional sample content ───────────────────────────────────────────
    if include_sample_content:
        doc.add_heading("1. Einleitung", level=1)
        p = doc.add_paragraph(
            "Dieses Dokument wurde mit dem Multiversum DOCX-Generator erstellt. "
            "Alle Stile, Kopf- und Fußzeilen sowie das Inhaltsverzeichnis entsprechen "
            "dem Multiversum Corporate Design."
        )
        _para_spacing(p)

        doc.add_heading("2. Hauptteil", level=1)
        doc.add_heading("2.1 Abschnitt", level=2)
        doc.add_paragraph(
            "Abschnittstext. Überschriften werden automatisch im Inhaltsverzeichnis "
            "erfasst (F9 in Word zum Aktualisieren)."
        )
        doc.add_heading("2.1.1 Unterabschnitt", level=3)
        doc.add_paragraph("Detailtext.")

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# CLI entry point
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate a Multiversum CI-compliant DOCX")
    parser.add_argument("--out", default="Multiversum_Dokument.docx", help="Output file path")
    parser.add_argument("--title", default="Dokumententitel", help="Document title")
    parser.add_argument("--subtitle", default="", help="Subtitle / description")
    parser.add_argument("--author", default="", help="Author name")
    parser.add_argument("--classification", default="intern",
                        choices=["öffentlich", "intern", "vertraulich", "streng vertraulich"])
    parser.add_argument("--valid-from", default=date.today().strftime("%d.%m.%Y"))
    parser.add_argument("--responsible", default="")
    parser.add_argument("--approved-by", default="")
    parser.add_argument("--approved-on", default="")
    parser.add_argument("--storage", default="")
    parser.add_argument("--sample", action="store_true", help="Include sample content sections")
    parser.add_argument("--no-toc", action="store_true", help="Skip table of contents")
    args = parser.parse_args()

    doc = create_mv_document(
        title=args.title,
        subtitle=args.subtitle,
        author=args.author,
        classification=args.classification,
        valid_from=args.valid_from,
        responsible=args.responsible,
        approved_by=args.approved_by,
        approved_on=args.approved_on,
        storage=args.storage,
        include_toc=not args.no_toc,
        include_sample_content=args.sample,
    )

    out_path = Path(args.out)
    doc.save(str(out_path))
    print(f"✓ Saved: {out_path.resolve()}")
    print(f"  Logo:   {'embedded' if LOGO_PATH.exists() else 'MISSING — check assets/'}")
    print(f"  TOC:    {'yes (press F9 in Word to update)' if not args.no_toc else 'skipped'}")


if __name__ == "__main__":
    main()
